import logging
import re
from collections.abc import Iterator
from pathlib import Path

from app.indexing.extractors._ooxml_guard import check_ooxml_archive

logger = logging.getLogger(__name__)

_WML_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_BOILERPLATE_NOTE_TYPES = {"separator", "continuationSeparator"}


def _note_text(note_el) -> str:
    """Join all `w:t` run text under a `w:footnote`/`w:endnote` element."""
    return "".join(t.text or "" for t in note_el.iter(f"{_WML_NS}t")).strip()


def _iter_headers_footers(doc):
    """Yield (label, _Header/_Footer) for each section's own (non-inherited)
    header/footer definitions.

    A linked section's `.header`/`.footer` resolves to the *prior* section's
    content, so processing it too would duplicate that text.
    """
    for section in doc.sections:
        for label, hf in (
            ("Header", section.header),
            ("First Page Header", section.first_page_header),
            ("Even Page Header", section.even_page_header),
            ("Footer", section.footer),
            ("First Page Footer", section.first_page_footer),
            ("Even Page Footer", section.even_page_footer),
        ):
            if not hf.is_linked_to_previous:
                yield label, hf


def _iter_footnotes_endnotes(doc):
    """Yield (label, text) for each real footnote/endnote in the package.

    python-docx doesn't expose footnotes/endnotes parts natively (they're
    absent from PartFactory.part_type_for), so they're reached directly via
    the OPC package relationships.
    """
    from docx.opc.constants import CONTENT_TYPE as CT
    from docx.oxml import parse_xml

    for part in doc.part.package.iter_parts():
        if part.content_type == CT.WML_FOOTNOTES:
            label, tag = "Footnote", f"{_WML_NS}footnote"
        elif part.content_type == CT.WML_ENDNOTES:
            label, tag = "Endnote", f"{_WML_NS}endnote"
        else:
            continue

        root = parse_xml(part.blob)
        for note_el in root.iter(tag):
            note_type = note_el.get(f"{_WML_NS}type")
            if note_type in _BOILERPLATE_NOTE_TYPES:
                continue
            text = _note_text(note_el)
            if text:
                yield label, text


_HEADING_RE = re.compile(r"^heading\s*(\d+)$", re.IGNORECASE)


def _heading_prefix(paragraph) -> str:
    """Markdown heading marker for a paragraph, or "" if it is body text.

    DOCX carries structure in paragraph styles, and dropping it flattened every
    document before chunking - a section title became indistinguishable from a
    sentence. Emitting markdown markers hands the chunker the same structural
    signal it already uses for .md, and the summarizer's markdown path builds
    its "Structure: A > B > C" summary from exactly these.
    """
    try:
        style = paragraph.style
        name = getattr(style, "name", None) if style is not None else None
        # isinstance, not truthiness: a malformed document can reference a style
        # missing from styles.xml, and python-docx substitutes are not always
        # strings. Feeding a non-string to the regex raises, and an exception
        # here would abort extraction of the entire document over one paragraph.
        if not isinstance(name, str) or not name.strip():
            return ""

        name = name.strip()
        if name.lower() == "title":
            return "# "
        match = _HEADING_RE.match(name)
        if not match:
            return ""
        # Word allows Heading 1-9; markdown stops at 6.
        level = min(int(match.group(1)), 6)
        return "#" * level + " "
    except Exception:
        logger.debug("Could not read paragraph style; treating as body text.", exc_info=True)
        return ""


class DocxExtractor:
    def can_handle(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def extract_stream(self, path: Path, max_file_size: int) -> Iterator[str]:
        """Yield text from paragraphs, tables, headers, footers, footnotes and
        endnotes in a DOCX document. Order: body, then headers/footers, then
        footnotes/endnotes."""
        try:
            check_ooxml_archive(path)

            from docx import Document

            doc = Document(str(path))
            total = 0

            from docx.table import Table
            from docx.text.paragraph import Paragraph

            for child in doc.element.body.iterchildren():
                if child.tag.endswith("}p"):
                    p = Paragraph(child, doc)
                    txt = p.text.strip()
                    if txt:
                        txt = _heading_prefix(p) + txt
                        yield txt
                        total += len(txt)
                        if total > max_file_size:
                            return
                elif child.tag.endswith("}tbl"):
                    t = Table(child, doc)
                    for row in t.rows:
                        row_data = [cell.text for cell in row.cells if cell.text.strip()]
                        if row_data:
                            line = " | ".join(row_data)
                            yield line
                            total += len(line)
                            if total > max_file_size:
                                return

            try:
                for label, hf in _iter_headers_footers(doc):
                    for p in hf.paragraphs:
                        txt = p.text.strip()
                        if txt:
                            yield f"[{label}] {txt}"
                            total += len(txt)
                            if total > max_file_size:
                                return
                    for t in hf.tables:
                        for row in t.rows:
                            row_data = [cell.text for cell in row.cells if cell.text.strip()]
                            if row_data:
                                line = " | ".join(row_data)
                                yield f"[{label}] {line}"
                                total += len(line)
                                if total > max_file_size:
                                    return
            except Exception as hf_err:
                logger.debug("Skipping unreadable header/footer in %s: %s", path, hf_err)

            try:
                for label, text in _iter_footnotes_endnotes(doc):
                    yield f"[{label}] {text}"
                    total += len(text)
                    if total > max_file_size:
                        return
            except Exception as note_err:
                logger.debug("Skipping unreadable footnotes/endnotes in %s: %s", path, note_err)
        except ValueError as ve:
            logger.warning("DOCX extraction aborted: %s", ve)
        except Exception as e:
            err_msg = str(e).lower()
            if "encrypted" in err_msg or "password" in err_msg:
                yield (
                    f"[ENCRYPTED DOCX: {path.name}] "
                    "Cannot extract text from password-protected file."
                )
                return
            logger.warning("Failed to extract DOCX %s: %s", path, e)

    def extract(self, path: Path, max_file_size: int) -> str:
        """Legacy extraction for backward compatibility."""
        return "\n".join(self.extract_stream(path, max_file_size))[:max_file_size]
